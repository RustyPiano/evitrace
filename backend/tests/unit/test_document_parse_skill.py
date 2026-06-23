from pathlib import Path

from app.config import settings
from app.models import Task, TaskFile
from app.services.task_service import serialize_file
from app.skills.base import SkillContext
from app.skills.document_parse import DocumentParseSkill, _hard_split, _merge_to_chunks, _paragraph_spans
from app.database import SessionLocal


def _task_file(path: Path, extension: str, modality: str = "document") -> TaskFile:
    task_id = path.parents[1].name
    with SessionLocal() as db:
        task = Task(name="Parse", objective="Objective", owner_id="owner", status="ready")
        task.id = task_id
        db.add(task)
        file = TaskFile(
            task_id=task_id,
            original_name=path.name,
            stored_name=path.name,
            extension=extension,
            mime_type=None,
            size_bytes=path.stat().st_size,
            modality=modality,
        )
        db.add(file)
        db.commit()
        db.refresh(file)
        return file


def _context(task_id: str) -> SkillContext:
    return SkillContext(task_id=task_id, run_id=None, data_root=str(settings.data_root_path))


def test_txt_parse_splits_on_blank_lines_and_skips_empty_blocks(tmp_path):
    task_id = "txt-task"
    original_dir = settings.data_root_path / "tasks" / task_id / "original"
    original_dir.mkdir(parents=True)
    path = original_dir / "note.txt"
    path.write_text("Alpha paragraph\n\n\nBeta paragraph\n\n", encoding="utf-8")
    file = _task_file(path, "txt", "text")

    result = DocumentParseSkill().run(_context(task_id), {"file": serialize_file(file)})

    assert result.success is True
    evidence = result.data["evidence"]
    assert len(evidence) == 1
    assert "Alpha paragraph" in evidence[0]["content"]
    assert "Beta paragraph" in evidence[0]["content"]
    assert evidence[0]["locator"] == {
        "kind": "text",
        "page": None,
        "paragraph": 1,
        "char_start": 0,
        "char_end": 32,
    }


def test_docx_parse_reads_non_empty_paragraphs(tmp_path):
    from docx import Document

    task_id = "docx-task"
    original_dir = settings.data_root_path / "tasks" / task_id / "original"
    original_dir.mkdir(parents=True)
    path = original_dir / "report.docx"
    document = Document()
    document.add_paragraph("First paragraph")
    document.add_paragraph("   ")
    document.add_paragraph("Second paragraph")
    document.save(path)
    file = _task_file(path, "docx")

    result = DocumentParseSkill().run(_context(task_id), {"file": serialize_file(file)})

    assert result.success is True
    evidence = result.data["evidence"]
    assert len(evidence) == 1
    assert "First paragraph" in evidence[0]["content"]
    assert "Second paragraph" in evidence[0]["content"]
    assert evidence[0]["locator"]["paragraph"] == 1


def test_pdf_parse_preserves_page_numbers(tmp_path):
    import fitz

    task_id = "pdf-task"
    original_dir = settings.data_root_path / "tasks" / task_id / "original"
    original_dir.mkdir(parents=True)
    path = original_dir / "report.pdf"
    pdf = fitz.open()
    page1 = pdf.new_page()
    page1.insert_text((72, 72), "First page text")
    page2 = pdf.new_page()
    page2.insert_text((72, 72), "Second page text")
    pdf.save(path)
    pdf.close()
    file = _task_file(path, "pdf")

    result = DocumentParseSkill().run(_context(task_id), {"file": serialize_file(file)})

    assert result.success is True
    evidence = result.data["evidence"]
    assert [item["locator"]["page"] for item in evidence] == [1, 2]
    assert evidence[0]["content"] == "First page text"
    assert evidence[1]["content"] == "Second page text"


def test_merge_to_chunks_combines_short_paragraphs_and_splits_at_target():
    text = "Alpha paragraph\n\nBeta paragraph\n\nGamma paragraph"
    spans = _paragraph_spans(text)

    merged = _merge_to_chunks(spans, target=100, max_chars=160)
    split = _merge_to_chunks(spans, target=31, max_chars=160)

    assert merged == [(text, 0, len(text))]
    assert [chunk[0] for chunk in split] == [
        "Alpha paragraph\n\nBeta paragraph",
        "Gamma paragraph",
    ]
    assert split[0][1:] == (0, 31)
    assert split[1][1:] == (33, 48)


def test_merge_with_source_content_equals_source_slice_and_respects_max_chars():
    # 原文段落间含多余空行/空白：content 必须等于 source[start:end]（locator 可还原 content）。
    text = "Alpha\n \n\nBeta\n\n\nGamma"
    spans = _paragraph_spans(text)

    chunks = _merge_to_chunks(spans, target=100, max_chars=200, source=text)

    assert len(chunks) == 1
    content, start, end = chunks[0]
    assert content == text[start:end]
    assert "Alpha" in content and "Beta" in content and "Gamma" in content

    # 大量短段落：合并块长度（span 跨度）绝不超过 max_chars，且仍合并（不是每段一条）。
    many = "\n\n".join(f"seg{index}" for index in range(60))
    many_chunks = _merge_to_chunks(_paragraph_spans(many), target=40, max_chars=80, source=many)
    assert len(many_chunks) < 60
    for content, start, end in many_chunks:
        assert content == many[start:end]
        assert end - start <= 80


def test_hard_split_prefers_boundaries_and_keeps_offsets_contiguous():
    text = "Alpha sentence. Beta sentence! Gamma sentence? Delta sentence."

    chunks = _hard_split(text, base_start=10, max_chars=30)

    assert len(chunks) > 1
    assert all(0 < len(content) <= 30 for content, _, _ in chunks)
    assert chunks[0][0].endswith("!")
    assert chunks[0][1:] == (10, 40)
    for content, start, end in chunks:
        assert content == text[start - 10 : end - 10]
    for previous, current in zip(chunks, chunks[1:]):
        assert previous[2] <= current[1]
        assert text[previous[2] - 10 : current[1] - 10].strip() == ""


def test_configured_chunk_sizes_affect_text_parse(monkeypatch, tmp_path):
    task_id = "config-text-task"
    original_dir = settings.data_root_path / "tasks" / task_id / "original"
    original_dir.mkdir(parents=True)
    path = original_dir / "note.txt"
    path.write_text("Alpha paragraph\n\nBeta paragraph\n\nGamma paragraph", encoding="utf-8")
    file = _task_file(path, "txt", "text")
    skill = DocumentParseSkill()

    monkeypatch.setattr(settings, "parse_chunk_target_chars", 100)
    monkeypatch.setattr(settings, "parse_chunk_max_chars", 160)
    merged = skill.run(_context(task_id), {"file": serialize_file(file)})

    monkeypatch.setattr(settings, "parse_chunk_target_chars", 31)
    monkeypatch.setattr(settings, "parse_chunk_max_chars", 160)
    split = skill.run(_context(task_id), {"file": serialize_file(file)})

    assert len(merged.data["evidence"]) == 1
    assert len(split.data["evidence"]) == 2


def test_scanned_pdf_without_text_returns_warning(tmp_path):
    import fitz

    task_id = "empty-pdf-task"
    original_dir = settings.data_root_path / "tasks" / task_id / "original"
    original_dir.mkdir(parents=True)
    path = original_dir / "empty.pdf"
    pdf = fitz.open()
    pdf.new_page()
    pdf.save(path)
    pdf.close()
    file = _task_file(path, "pdf")

    result = DocumentParseSkill().run(_context(task_id), {"file": serialize_file(file)})

    assert result.success is True
    assert result.data["evidence"] == []
    assert result.warnings == ["PDF 未提取到文本，可能是扫描件"]
